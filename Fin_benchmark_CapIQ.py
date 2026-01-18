import os
import sys
import time
import socket
import subprocess
import webbrowser
from typing import List, Tuple, Dict, Optional
import math 
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st
import json
import io
import re
from io import BytesIO
from difflib import SequenceMatcher
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.platypus import Image as RLImage, PageBreak  
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase.pdfmetrics import stringWidth

# =============================================================================
# Self-launching Streamlit bootstrap
# =============================================================================
def _running_inside_streamlit() -> bool:
    try:
        from streamlit.runtime.scriptrunner import get_script_run_ctx
        if get_script_run_ctx() is not None:
            return True
    except Exception:
        pass
    try:
        from streamlit import runtime
        if runtime.exists():
            return True
    except Exception:
        pass
    return os.environ.get("STREAMLIT_RUN") == "1"


def _wait_for_port(host: str, port: int, timeout: float = 30.0) -> bool:
    start = time.time()
    while time.time() - start < timeout:
        try:
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
                sock.settimeout(0.5)
                sock.connect((host, int(port)))
                return True
        except OSError:
            time.sleep(0.3)
    return False


def _spawn_streamlit():
    """Run `streamlit run` on this file when launched with python."""
    if _running_inside_streamlit():
        main()  # IMPORTANT: call the real UI
        return

    port = os.environ.get("PORT", "8501")
    env = os.environ.copy()
    env["STREAMLIT_RUN"] = "1"
    cmd = [
        sys.executable, "-m", "streamlit", "run", os.path.abspath(__file__),
        "--server.port", port, "--server.headless", "true",
        "--browser.gatherUsageStats=false",
    ]
    try:
        print(f"[INFO] Starting Streamlit on http://localhost:{port} ...")
        proc = subprocess.Popen(cmd, env=env)
        if _wait_for_port("localhost", int(port), timeout=30.0):
            try:
                webbrowser.open_new_tab(f"http://localhost:{port}")
            except Exception:
                pass
        proc.wait()
    except FileNotFoundError:
        print("[WARN] streamlit not found; running inline.")
        os.environ["STREAMLIT_RUN"] = "1"
        main()


# =============================================================================
# App constants
# =============================================================================
DATA_FILE = "Company_Financials_CapIQ.xlsx"
DATA_SHEET = "data"
METRICS_TYPE_SHEET = "metrics_type"

AUDIT_DB_FILE = "Audit_Work_Program.xlsx"
AUDIT_DB_SHEET = None

CURRENCY_BY_EXCHANGE = {
    "SGX": "SGD", "NYSE": "USD", "NYSE ARCA": "USD", "NYSE MKT": "USD",
    "BATS": "USD", "LSE": "GBP", "HKEX": "HKD", "Catalist": "SGD"
}

CHIP_CSS = """
<style>
.chip {display:inline-block; padding:3px 10px; border-radius:999px; font-size:12px; font-weight:700; margin-right:6px;}
.chip-green {background:#def7e5; color:#065f46; border:1px solid #34d399;}
.chip-amber {background:#fff3cd; color:#8a6d3b; border:1px solid #fbbf24;}
.chip-red   {background:#fde2e4; color:#7f1d1d; border:1px solid #f87171;}
.kpi {border:1px solid #e5e7eb; border-radius:8px; padding:10px 12px; margin:6px 0; background:#fff;}
.hr {height:1px; background:#eee; border:none; margin:10px 0 6px 0;}
.legend {font-size:12px; color:#4b5563; margin: 4px 0 10px 0;}
</style>
"""


# =============================================================================
# Data helpers
# =============================================================================
@st.cache_data(show_spinner=False)
def load_all_sheets(path: str = DATA_FILE) -> Dict[str, pd.DataFrame]:
    if not os.path.exists(path):
        st.error(f"Missing Excel file: {path} (expected in the current folder).")
        st.stop()
    xl = pd.ExcelFile(path, engine="openpyxl")
    out = {
        DATA_SHEET: xl.parse(DATA_SHEET),
        METRICS_TYPE_SHEET: xl.parse(METRICS_TYPE_SHEET),
    }
    return out


def infer_metric_columns(df: pd.DataFrame) -> List[str]:
    cols = list(df.columns)
    if "FY" not in cols:
        raise ValueError("Expected 'FY' column in data sheet")
    # All columns to the right of FY are interpreted as metric columns
    return cols[cols.index("FY") + 1 :]


def to_numeric(df: pd.DataFrame, cols: List[str]) -> pd.DataFrame:
    out = df.copy()
    for c in cols:
        out[c] = pd.to_numeric(out[c], errors="coerce")
    return out

@st.cache_data(show_spinner=False)
def compute_percentiles(
    df: pd.DataFrame,
    metric_cols: List[str],
    group_cols: Tuple[str, str, str] = ("EXCHANGE", "INDUSTRY", "FY"),
) -> Tuple[pd.DataFrame, pd.DataFrame]:
    """
    Returns:
      - wide percentiles (MultiIndex columns: (metric, stat))
      - tidy long dataframe with columns [EXCHANGE, INDUSTRY, FY, metric, p25, p50, p75]
    """
    df_num = df[list(group_cols) + list(metric_cols)].copy()
    for m in metric_cols:
        df_num[m] = pd.to_numeric(df_num[m], errors="coerce")

    q = df_num.groupby(list(group_cols))[metric_cols].quantile([0.25, 0.5, 0.75]).unstack(-1)
    new_cols = []
    for metric, qlevel in q.columns:
        tag = {0.25: "p25", 0.5: "p50", 0.75: "p75"}.get(float(qlevel), str(qlevel))
        new_cols.append((metric, tag))
    q.columns = pd.MultiIndex.from_tuples(new_cols, names=["metric", "stat"])

    # Build a tidy view
    rows = []
    for m in metric_cols:
        sub = q[m].reset_index()
        sub.insert(3, "metric", m)
        rows.append(sub)
    tidy = pd.concat(rows, ignore_index=True)
    return q, tidy

@st.cache_data(show_spinner=False)
def slice_company_row_for_fy(df, exchange, industry, fy, company) -> pd.DataFrame:
    mask = (
        (df["EXCHANGE"].astype(str) == exchange) &
        (df["INDUSTRY"].astype(str) == industry) &
        (df["FY"].astype(str) == str(fy))
    )
    df_slice = df.loc[mask]
    if company:
        exact = df_slice[df_slice["ENTITY_NAME"].astype(str).str.strip().str.lower() == company.strip().lower()]
        if not exact.empty:
            return exact.iloc[[0]]
    return df_slice.iloc[[0]] if not df_slice.empty else pd.DataFrame()

@st.cache_data(show_spinner=False)
def yoy_company_series(df, exchange, industry, company) -> pd.DataFrame:
    mask = (
        (df["EXCHANGE"].astype(str) == exchange) &
        (df["INDUSTRY"].astype(str) == industry) &
        (df["ENTITY_NAME"].astype(str).str.strip().str.lower() == company.strip().lower())
    )
    out = df.loc[mask].copy()
    out["FY"] = out["FY"].astype(str)
    return out

@st.cache_data(show_spinner=False)
def load_audit_db(path: str = AUDIT_DB_FILE, sheet: str | None = AUDIT_DB_SHEET) -> pd.DataFrame:
    if not os.path.exists(path):
        st.error(f"Missing audit master Excel file: {path}")
        st.stop()
    xl = pd.ExcelFile(path, engine="openpyxl")
    sh = sheet or xl.sheet_names[0]
    df = xl.parse(sh)

    # Normalize expected column names (exact names in your file)
    expected = ["Scope", "Sub-process", "Risk", "Control Description", "Audit Test Steps", "Documents required"]
    missing = [c for c in expected if c not in df.columns]
    if missing:
        st.error(f"Audit DB missing columns: {missing} (expected: {expected})")
        st.stop()

    # Basic cleanup
    for c in expected:
        df[c] = df[c].astype(str).str.strip()
    return df

@st.cache_data(show_spinner=False)
def audit_vocab(df: pd.DataFrame) -> tuple[list[str], dict[str, list[str]]]:
    scopes = sorted(df["Scope"].dropna().unique().tolist())
    subs_by_scope = {s: sorted(df.loc[df["Scope"] == s, "Sub-process"].dropna().unique().tolist()) for s in scopes}
    return scopes, subs_by_scope

# ==============================
# Evidence extraction utilities
# ==============================
try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

MAX_FILE_CHARS = 60_000          # hard cap per file after parsing
CHUNK_SIZE_CHARS = 2_000         # size for scoring chunks
TOP_K_CHUNKS_PER_FILE = 5        # keep top-k chunks per file
MAX_TOTAL_EXCERPTS = 20          # cross-file cap to limit tokens

def _norm(s: str) -> str:
    s = str(s or "")
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _score_overlap(query: str, text: str) -> float:
    """Simple blended score using sequence similarity and token overlap."""
    from math import sqrt
    q = _norm(query).lower()
    t = _norm(text).lower()
    if not q or not t:
        return 0.0
    seq = SequenceMatcher(None, q, t).ratio()
    q_tokens = set(re.findall(r"[a-z0-9]+", q))
    t_tokens = set(re.findall(r"[a-z0-9]+", t))
    if not q_tokens or not t_tokens:
        return seq
    jacc = len(q_tokens & t_tokens) / max(1, len(q_tokens | t_tokens))
    return 0.6 * seq + 0.4 * jacc

def _chunk_text(text: str, size: int = CHUNK_SIZE_CHARS) -> list[str]:
    text = _norm(text)
    if len(text) <= size:
        return [text]
    chunks = [text[i:i+size] for i in range(0, min(len(text), MAX_FILE_CHARS), size)]
    return chunks

def extract_text_from_upload(uploaded_file) -> tuple[str, dict]:
    name = getattr(uploaded_file, "name", "file")
    meta = {"name": name, "size": getattr(uploaded_file, "size", None), "type": ""}
    raw = uploaded_file.read()
    uploaded_file.seek(0)  # reset

    lower = name.lower()
    text = ""
    try:
        if lower.endswith(".pdf") and PyPDF2 is not None:
            meta["type"] = "pdf"
            reader = PyPDF2.PdfReader(io.BytesIO(raw))
            pages = []
            for i, p in enumerate(reader.pages):
                try:
                    pages.append(p.extract_text() or "")
                except Exception:
                    pages.append("")
            text = "\n\n".join([f"[Page {i+1}] {t}" for i, t in enumerate(pages)])
        elif lower.endswith(".docx") and docx is not None:
            meta["type"] = "docx"
            d = docx.Document(io.BytesIO(raw))
            paras = [p.text for p in d.paragraphs]
            text = "\n".join(paras)
            # tables (simple)
            for tbl in d.tables:
                for r in tbl.rows:
                    cells = [c.text for c in r.cells]
                    if any(_norm(x) for x in cells):
                        text += "\n" + " | ".join(cells)
        elif lower.endswith(".xlsx") or lower.endswith(".xls"):
            meta["type"] = "excel"
            engine = "openpyxl" if lower.endswith(".xlsx") else "xlrd"
            xls = pd.ExcelFile(io.BytesIO(raw), engine=engine)
            parts = []
            for si, sh in enumerate(xls.sheet_names[:3]):
                try:
                    df = xls.parse(sh)
                except Exception:
                    continue
                if df.empty:
                    continue
                head = df.head(15).fillna("").astype(str)
                sample_txt = head.to_csv(index=False)
                parts.append(f"[Sheet {si+1}: {sh}]\n{sample_txt}")
            text = "\n\n".join(parts)
        elif lower.endswith(".csv"):
            meta["type"] = "csv"
            df = pd.read_csv(io.BytesIO(raw))
            if not df.empty:
                text = df.head(50).to_csv(index=False)
        elif lower.endswith(".txt"):
            meta["type"] = "txt"
            text = raw.decode("utf-8", errors="ignore")
        else:
            meta["type"] = "binary"
            text = ""
    except Exception as e:
        meta["error"] = f"parse_error: {e}"
        text = ""

    if len(text) > MAX_FILE_CHARS:
        text = text[:MAX_FILE_CHARS] + "\n...[truncated]..."
    return text, meta

def build_relevant_excerpts(test_steps: str, docs: list[dict]) -> list[dict]:
    q = _norm(test_steps)
    all_excerpts = []
    for d in docs:
        chunks = _chunk_text(d.get("text", ""))
        scored = [(c, _score_overlap(q, c)) for c in chunks]
        scored.sort(key=lambda x: x[1], reverse=True)
        for c, sc in scored[:TOP_K_CHUNKS_PER_FILE]:
            all_excerpts.append({
                "name": d["meta"].get("name"),
                "type": d["meta"].get("type"),
                "excerpt": c,
                "score": float(sc),
            })
    all_excerpts.sort(key=lambda x: x["score"], reverse=True)
    return all_excerpts[:MAX_TOTAL_EXCERPTS]

def _parse_documents_required(cell_value: str) -> list[str]:
    """
    Turn the 'Documents required' cell (Excel column F) into clean labels.
    Supports bullets, dashes, semicolons, and numbered lists.
    """
    raw = str(cell_value or "")
    # Normalize common bullets/delimiters to newline
    raw = raw.replace("•", "\n").replace("‣", "\n").replace("◦", "\n")
    # Split on newlines and semicolons as fallback
    parts = []
    for line in re.split(r"[\n;]+", raw):
        # Trim typical bullets/numbers like '-', '–', '—', '*', '1.', 'a)', etc.
        line = re.sub(r"^\s*[-–—*•\d\.\)]\s*", "", str(line)).strip()
        if line:
            parts.append(line)
    # Deduplicate while preserving order
    seen, out = set(), []
    for p in parts:
        if p not in seen:
            seen.add(p)
            out.append(p)
    return out

# =============================================================================
# Matching & autopopulate helpers 
# =============================================================================

def _safe_take(df: pd.DataFrame, cols: list[str]) -> tuple[pd.DataFrame, list[str]]:
    cols = list(cols)
    missing = [c for c in cols if c not in df.columns]
    if missing:
        return pd.DataFrame(), missing
    return df[cols].copy(), []

def get_company_matches(df: pd.DataFrame, typed: str) -> List[str]:
    if not typed or not str(typed).strip():
        return []
    names = df["ENTITY_NAME"].dropna().astype(str).unique().tolist()
    typed_low = str(typed).strip().lower()
    return [n for n in names if typed_low in n.lower()][:10]


def find_exact_company(df: pd.DataFrame, typed: str) -> pd.DataFrame:
    """Return all rows for the exact case-insensitive company name, else empty df."""
    mask = df["ENTITY_NAME"].astype(str).str.strip().str.lower() == str(typed).strip().lower()
    return df.loc[mask].copy()


def latest_profile_rows(df_company: pd.DataFrame) -> pd.DataFrame:
    """Return rows sorted to have the latest FY first."""
    if df_company.empty:
        return df_company
    df_company = df_company.copy()
    df_company["__FY_"] = pd.to_numeric(df_company["FY"], errors="coerce")
    return df_company.sort_values("__FY_", ascending=False).drop(columns="__FY_", errors="ignore")


def pick_best_profile(df_company: pd.DataFrame) -> Dict[str, str]:
    """Pick the (EXCHANGE, INDUSTRY) from the latest FY row for the company."""
    if df_company.empty:
        return {}
    df_latest = latest_profile_rows(df_company)
    row = df_latest.iloc[0]
    return {"EXCHANGE": str(row.get("EXCHANGE", "")), "INDUSTRY": str(row.get("INDUSTRY", ""))}


def _first_present_column(df: pd.DataFrame, aliases: List[str]) -> str:
    """Find the first column present in df (case-insensitive) from aliases."""
    cols_low = {c.lower(): c for c in df.columns}
    for a in aliases:
        if a.lower() in cols_low:
            return cols_low[a.lower()]
    return ""


def read_key_financials_for(df_company: pd.DataFrame, prefer_fy: str) -> Dict[str, float]:
    if df_company.empty:
        return {}

    dfc = df_company.copy()
    dfc["FY_str"] = dfc["FY"].astype(str)
    if prefer_fy and (dfc["FY_str"] == str(prefer_fy)).any():
        row = dfc.loc[dfc["FY_str"] == str(prefer_fy)].iloc[0]
    else:
        row = latest_profile_rows(dfc).iloc[0]

    # --- Aliases (case-insensitive) updated for your CapIQ columns ---
    # Feel free to extend these lists if you discover more CapIQ names.
    alias_map = {
        "Current Assets":        ["TOTAL_CA", "CURRENT_ASSETS", "Current Assets"],
        "Current Liabilities":   ["TOTAL_CL", "CURRENT_LIABILITIES", "Current Liabilities"],
        "Inventory":             ["INVENTORY", "Inventory"],
        "Operating Cash Flow":   ["TOTAL_OPER_EXPEN", "OPERATING_CASH_FLOW", "Operating Cash Flow", "CFO", "NET_CASH_FROM_OPERATIONS"],
        "Capital Expenditure":   ["CAPEX", "Capital Expenditure", "CAPITAL_EXPENDITURE"],
        "Revenue":               ["TOTAL_REV", "REVENUE", "Revenue", "TOTAL_REVENUE", "SALES"],
        "EBITDA":                ["EBITDA", "Ebitda"],
        "Cost of Revenue":       ["COST_OF_REVENUE", "Cost of Revenue", "COGS"],
        # Optional: surface market cap, kept separate b/c units differ
        "MARKETCAP ($'M)":       ["MARKETCAP ($'M)", "MarketCap_Millions", "MARKETCAP_M"]
    }

    # helper: first matching column name in df (case-insensitive)
    cols_low = {c.lower(): c for c in dfc.columns}
    def pick_col(aliases: list[str]) -> str:
        for a in aliases:
            if a.lower() in cols_low:
                return cols_low[a.lower()]
        return ""

    # Read values
    raw = {}
    for label, aliases in alias_map.items():
        col = pick_col(aliases)
        if not col:
            continue
        v = pd.to_numeric(row.get(col, np.nan), errors="coerce")
        if pd.notna(v):
            raw[label] = float(v)

    # --- Scale: thousands -> units for all standard financials; market cap stays millions ---
    OUT = {}
    for label, v in raw.items():
        if label == "MARKETCAP ($'M)":
            # keep in millions; if you prefer units: OUT["Market Cap"] = v * 1_000_000
            OUT[label] = v
        else:
            OUT[label] = v * 1_000.0    # scale k -> units

    return OUT


def _try_get_percentile_row(pct_wide: pd.DataFrame, exch: str, ind: str, fy: str):
    """Robust access to percentile row by attempting string and numeric FY keys."""
    if pct_wide is None:
        return None
    try:
        return pct_wide.loc[(exch, ind, fy)]
    except Exception:
        pass
    try:
        fy_num = float(fy) if fy is not None else None
        return pct_wide.loc[(exch, ind, fy_num)]
    except Exception:
        return None

# Clear Tab 1 analysis artifacts
def _clear_tab1_analysis():
    st.session_state.pop("fin_analysis_text", None)
    st.session_state.pop("fin_analysis_pdf_bytes", None)

# Convert simple Markdown text to PDF bytes 

    def md_to_pdf_bytes(
        md_text: str,
        title: str = "",
        author: str = "",
        subheaders: Optional[Dict[str, str]] = None,   # {"Company": "...", "Exchange": "...", "Industry": "...", "FY": "..."}
    ) -> bytes:

        subheaders = subheaders or {}

        # --- Sanitize text: remove non-printable/control chars that sometimes sneak in as ''
        def _sanitize(s: str) -> str:
            if not s:
                return ""
            # Keep common whitespace + printable; drop control chars except \n and \t
            return re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", str(s))

        md_text = _sanitize(md_text)
        clean_sub = { _sanitize(k): _sanitize(v) for k, v in subheaders.items() if k and v }

        # --- Layout constants
        PAGE_W, PAGE_H = A4
        # Header block we draw above the story frame:
        HEADER_H = 24 * mm     # header band height (title + subheaders + divider)
        BODY_TOP_PAD = 6 * mm  # cushion between header divider and first body line

        # The story's top margin is header space + cushion + some normal top margin
        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            title=title or "Report",
            author=author or "",
            leftMargin=16 * mm,
            rightMargin=16 * mm,
            topMargin=HEADER_H + BODY_TOP_PAD,      # <<— ensures the frame starts BELOW our header
            bottomMargin=18 * mm,
        )

        styles = getSampleStyleSheet()
        body = styles["BodyText"]; body.leading = 14
        h2 = styles["Heading2"];   h3 = styles["Heading3"]

        # --- Build story from simple Markdown-ish input (same behavior you had)
        elements = []
        bullets_acc = []

        def flush_bullets():
            nonlocal bullets_acc, elements
            if bullets_acc:
                items = [ListItem(Paragraph(b, body)) for b in bullets_acc]
                elements.append(ListFlowable(items, bulletType="bullet", start="•"))
                elements.append(Spacer(1, 6))
                bullets_acc = []

        for raw in (md_text or "").splitlines():
            line = raw.rstrip()

            if not line.strip():
                flush_bullets()
                elements.append(Spacer(1, 6))
                continue

            if line.startswith("### "):
                flush_bullets()
                elements.append(Paragraph(_sanitize(line[4:].strip()), h3))
                continue
            if line.startswith("## "):
                flush_bullets()
                elements.append(Paragraph(_sanitize(line[3:].strip()), h2))
                continue
            if line.startswith("# "):
                flush_bullets()
                elements.append(Paragraph(_sanitize(line[2:].strip()), h2))
                continue

            if line.lstrip().startswith("- ") or line.lstrip().startswith("* "):
                bullets_acc.append(_sanitize(line.lstrip()[2:].strip()))
                continue

            # Simple **bold** handling
            htmlish = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", _sanitize(line))
            elements.append(Paragraph(htmlish, body))

        flush_bullets()

        # --- Header painter (drawn completely above story frame)
        def _split_line_to_width(text: str, font_name: str, font_size: float, max_width: float) -> list[str]:
            """Greedy wrap into multiple lines based on measured width."""
            words = text.split()
            if not words:
                return []
            lines, cur = [], words[0]
            for w in words[1:]:
                trial = cur + " " + w
                if stringWidth(trial, font_name, font_size) <= max_width:
                    cur = trial
                else:
                    lines.append(cur)
                    cur = w
            lines.append(cur)
            return lines

        def _draw_header(c: canvas.Canvas, doc_obj):
            x_left = doc.leftMargin
            x_right = PAGE_W - doc.rightMargin
            usable_w = x_right - x_left

            # The header occupies the band [PAGE_H - HEADER_H, PAGE_H], entirely ABOVE the story frame.
            y_top = PAGE_H - (HEADER_H * 0.55)  # small visual offset inside the band

            # Title
            title_text = _sanitize(title or "Report")
            c.setFont("Helvetica-Bold", 14)
            c.setFillColor(colors.black)
            c.drawString(x_left, y_top, title_text)

            # Subheaders as "Label: value" separated by •, wrapped safely
            c.setFont("Helvetica", 9)
            c.setFillColor(colors.HexColor("#374151"))
            pairs = [f"{k}: {v}" for k, v in clean_sub.items() if k and v]
            y = y_top - 16  # gap after title

            if pairs:
                sep = "   •   "
                full = sep.join(pairs)
                # If too wide, wrap across multiple rows
                for line in _split_line_to_width(full, "Helvetica", 9, usable_w):
                    c.drawString(x_left, y, line)
                    y -= 12

            # Divider line at the bottom edge of the header band
            c.setStrokeColor(colors.HexColor("#e5e7eb"))
            c.setLineWidth(0.8)
            y_div = PAGE_H - HEADER_H  # bottom of header band
            c.line(x_left, y_div, x_right, y_div)

            # Optional page number footer
            c.setFont("Helvetica", 8)
            c.setFillColor(colors.HexColor("#6b7280"))
            c.drawRightString(x_right, 12 * mm, f"Page {doc.page}")

        # Hook: draw the same header on all pages
        def _on_any_page(c, doc_obj):
            _draw_header(c, doc_obj)

        # Build
        doc.build(elements, onFirstPage=_on_any_page, onLaterPages=_on_any_page)
        return buf.getvalue()

def analysis_and_charts_to_html_bytes(
    analysis_md_text: str,
    figs: list[tuple[str, str, "go.Figure"]],
    title: str = "",
    subtitle: str = "",
    skipped: list[str] | None = None,
) -> bytes:
    # Minimal page shell + reuse your chip CSS
    css = """
    <style>
      body { font-family: system-ui, -apple-system, Segoe UI, Roboto, Arial, sans-serif; margin: 18px; }
      h1,h2,h3 { margin: 0.2em 0 0.2em; }
      .legend { font-size: 12px; color: #4b5563; margin: 6px 0 12px; }
      .hr { height:1px; background:#eee; border:none; margin:10px 0 6px 0; }
      .chip {display:inline-block;padding:3px 10px;border-radius:999px;font-size:12px;font-weight:700;margin-right:6px;}
      .chip-green {background:#def7e5;color:#065f46;border:1px solid #34d399;}
      .chip-amber {background:#fff3cd;color:#8a6d3b;border:1px solid #fbbf24;}
      .chip-red {background:#fde2e4;color:#7f1d1d;border:1px solid #f87171;}
      .section { margin: 18px 0; }
      .muted { color:#6b7280; }
      .box { border:1px solid #e5e7eb; border-radius:8px; padding:10px 12px; background:#fff; }
      .mt8 { margin-top: 8px; }
    </style>
    """
    parts = [f"<!DOCTYPE html><html><head><meta charset='utf-8'>{css}</head><body>"]
    if title:
        parts.append(f"<h2>{title}</h2>")
    if subtitle:
        parts.append(f"<div class='muted'>{subtitle}</div>")

     # Charts (load Plotly.js once via CDN on the first figure)
    parts.append("<div class='section'><h3>Benchmarking Charts</h3>")
    first = True
    for i, (metric_name, bucket, fig) in enumerate(figs, start=1):
        chip_class = {"Healthy":"chip-green","Satisfactory":"chip-amber","Needs Improvement":"chip-red"}.get(bucket, "chip-amber")
        parts.append(f"<div class='mt8'><div><strong>{i}. {metric_name}</strong> — <span class='chip {chip_class}'>{bucket}</span></div>")
        parts.append(fig.to_html(full_html=False, include_plotlyjs='cdn' if first else False, config={"displaylogo": False}))
        parts.append("</div>")
        first = False
    parts.append("</div>")  # end charts

    if skipped:
        parts.append("<div class='section'><h3>Skipped metrics</h3>")
        parts.append(f"<div class='muted'>{', '.join(skipped)}</div></div>")
 
    # Analysis block (keep simple; browsers render Markdown-ish text fine)
    if analysis_md_text and analysis_md_text.strip():
        parts.append("<div class='section'><h3>Financial Metrics – Analysis</h3>")
        # escape minimal HTML; or render Markdown if you have 'markdown' installed
        safe = analysis_md_text.replace("<", "&lt;").replace(">", "&gt;")
        # lightweight formatting for bullets/headers:
        safe = safe.replace("\n- ", "<br>• ").replace("\n* ", "<br>• ")
        parts.append(f"<div class='box mt8' style='white-space:pre-wrap'>{safe}</div></div>")

    parts.append("</body></html>")
    return "\n".join(parts).encode("utf-8")

# =============================================================================
# Visuals
# =============================================================================
def _classify_bucket(value: float, p25: float, p50: float, p75: float, grade: str) -> str:
    import pandas as _pd
    if _pd.isna(value) or _pd.isna(p25) or _pd.isna(p50) or _pd.isna(p75):
        return "—"
    lower = str(grade).lower().strip() == "lower"  # lower is better
    if lower:
        return "Healthy" if value <= p25 else ("Needs Improvement" if value >= p75 else "Satisfactory")
    else:
        return "Healthy" if value >= p75 else ("Needs Improvement" if value <= p25 else "Satisfactory")


def _chip_class(bucket: str) -> str:
    return {"Healthy": "chip-green", "Satisfactory": "chip-amber", "Needs Improvement": "chip-red"}.get(bucket, "chip-amber")

def _plot_benchmark(
    value, p25, p50, p75,
    grade,
    title="",
    company_name="Company",
    band_thickness=0.16,
    qline_width=1.2,
    cross_size=18,
    # white pill label by default
    label_bgcolor="#FFFFFF",
    label_bordercolor="#1d4ed8",
    label_font_color="#1d4ed8",
    domain_mode: str = "auto",   # "auto" | "minmax" | "whisker"
    pad_ratio: float = 0.08      # ~8% padding around chosen domain
):
   
    import pandas as _pd
    import numpy as _np
    import math
    import plotly.graph_objects as go

    fig = go.Figure()

    # --- normalize inputs
    def f(x):
        try:
            return float(x)
        except Exception:
            return float("nan")

    v, q25, q50, q75 = f(value), f(p25), f(p50), f(p75)
    xs = [x for x in [q25, q50, q75, v] if _pd.notna(x) and math.isfinite(x)]
    if not xs:
        xs = [0.0, 1.0]  # fallback

    # --- choose domain so it always includes the company value and percentiles
    # modes:
    #  - "minmax": [min(all), max(all)]
    #  - "whisker": use Tukey whiskers if IQR available; otherwise minmax
    #  - "auto": start with minmax, then expand to whiskers if quartiles exist
    def have_quartiles():
        return _pd.notna(q25) and _pd.notna(q75) and (q75 >= q25)

    min_x = min(xs)
    max_x = max(xs)

    if domain_mode == "minmax":
        pass  # already min/max of available values
    elif domain_mode == "whisker":
        if have_quartiles():
            iqr = q75 - q25
            min_x = min(min_x, q25 - 1.5 * iqr)
            max_x = max(max_x, q75 + 1.5 * iqr)
    else:  # "auto" (default): minmax, expanded by whiskers if present
        if have_quartiles():
            iqr = q75 - q25
            min_x = min(min_x, q25 - 1.5 * iqr)
            max_x = max(max_x, q75 + 1.5 * iqr)

    # guard against degenerate span
    span = max(1e-12, (max_x - min_x))
    pr = max(0.0, float(pad_ratio))
    x0 = min_x - pr * span
    x1 = max_x + pr * span

    # --- slim band geometry
    y_mid = 0.50
    h = max(0.06, min(0.35, float(band_thickness)))
    y0, y1 = y_mid - h, y_mid + h

    # --- background bands (drawn BELOW traces)
    lower_is_better = str(grade).lower().strip() == "lower"

    def rect(a, b, color):
        a, b = (a, b) if a <= b else (b, a)
        fig.add_shape(
            type="rect", x0=a, x1=b, y0=y0, y1=y1,
            fillcolor=color, line_width=0, layer="below"
        )

    # default positions if quartiles missing: split the domain roughly into thirds
    _q25 = q25 if _pd.notna(q25) else (x0 + (x1 - x0) * 0.33)
    _q75 = q75 if _pd.notna(q75) else (x0 + (x1 - x0) * 0.66)

    if lower_is_better:
        rect(x0, _q25, "#def7e5")   # healthy (lower)
        rect(_q25, _q75, "#fff3cd") # satisfactory
        rect(_q75, x1, "#fde2e4")   # needs improvement
    else:
        rect(x0, _q25, "#fde2e4")   # needs improvement
        rect(_q25, _q75, "#fff3cd") # satisfactory
        rect(_q75, x1, "#def7e5")   # healthy

    # --- quartile ticks (also BELOW traces)
    def tick(x, color, width):
        if _pd.notna(x):
            fig.add_shape(
                type="line", x0=x, x1=x, y0=y0, y1=y1,
                line=dict(color=color, width=width),
                layer="below"
            )

    tick(q25, "#9ca3af", qline_width)
    tick(q50, "#6b7280", qline_width)
    tick(q75, "#374151", qline_width)

    # --- quartile labels (annotations are above)
    label_y = y1 + 0.05

    def qlabel(x, txt, color):
        if _pd.notna(x):
            fig.add_annotation(
                x=x, y=label_y, text=txt, xanchor="center", yanchor="bottom",
                showarrow=False, font=dict(color=color, size=10)
            )

    qlabel(q25, "p25", "#9ca3af")
    qlabel(q50, "p50", "#6b7280")
    qlabel(q75, "p75", "#374151")

    # --- company marker (on TOP of shapes)
    if _pd.notna(v):
        fig.add_trace(
            go.Scatter(
                x=[v], y=[y_mid],
                mode="markers",
                marker=dict(symbol="x", size=cross_size, color="#1d4ed8"),
                name="Company",
                hovertemplate=f"{company_name}<extra></extra>",
                cliponaxis=False  # keep marker visible at the edges
            )
        )

        # white pill label just below the band
        # (clamp within axis so it never disappears)
        def _clamp(val, lo, hi):
            try:
                return max(lo, min(hi, float(val)))
            except Exception:
                return (lo + hi) * 0.5

        label_x = _clamp(v, x0, x1)
        fig.add_annotation(
            x=label_x, y=y0 - 0.06,
            text=company_name or "Company",
            xanchor="center", yanchor="top",
            showarrow=False,
            bgcolor=label_bgcolor,
            bordercolor=label_bordercolor,
            borderwidth=1,
            borderpad=6,
            font=dict(color=label_font_color, size=12)
        )

    # --- axes & layout
    fig.update_xaxes(
        range=[x0, x1],
        showgrid=True, gridcolor="#e5e7eb",
        zeroline=False, tickmode="auto"
    )
    fig.update_yaxes(visible=False, range=[0, 1])
    fig.update_layout(
        height=110,
        margin=dict(l=6, r=6, t=6, b=12),
        plot_bgcolor="#ffffff",
        paper_bgcolor="#ffffff",
        title=title,
        showlegend=False
    )
    return fig


def _plot_yoy(industry_df: pd.DataFrame, company_df: pd.DataFrame, metric: str, grade: str) -> go.Figure:
    fig = go.Figure()
    if industry_df is not None and not industry_df.empty:
        lower = str(grade).lower().strip() == "lower"
        if lower:
            fig.add_trace(go.Scatter(x=industry_df["FY"], y=industry_df["p25"], name="p25", line=dict(color="green")))
            fig.add_trace(go.Scatter(x=industry_df["FY"], y=industry_df["p50"], name="p50",
                                     line=dict(color="orange", dash="dot")))
            fig.add_trace(go.Scatter(x=industry_df["FY"], y=industry_df["p75"], name="p75", line=dict(color="red")))
        else:
            fig.add_trace(go.Scatter(x=industry_df["FY"], y=industry_df["p25"], name="p25", line=dict(color="red")))
            fig.add_trace(go.Scatter(x=industry_df["FY"], y=industry_df["p50"], name="p50",
                                     line=dict(color="orange", dash="dot")))
            fig.add_trace(go.Scatter(x=industry_df["FY"], y=industry_df["p75"], name="p75", line=dict(color="green")))
    if company_df is not None and not company_df.empty:
        fig.add_trace(go.Scatter(
            x=company_df["FY"], y=company_df[metric], name="Company",
            mode="lines+markers", line=dict(color="blue", width=3)
        ))
    fig.update_layout(height=220, legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                      xaxis_title="FY", yaxis_title=metric)
    return fig


# =========================
# YoY helpers
# =========================


def _fy_to_num(x) -> Optional[float]:
    s = str(x or "").strip()
    if s == "":
        return None

    # 1) 4-digit year anywhere in the string
    m4 = re.search(r"(19|20)\d{2}", s)
    if m4:
        return float(m4.group(0))

    # 2) 2-digit patterns like 'FY21', '21A' → map to 20xx for 00–79; 19xx for 80–99
    m2 = re.search(r"(?<!\d)(\d{2})(?!\d)", s)
    if m2:
        yy = int(m2.group(1))
        return float(2000 + yy) if yy <= 79 else float(1900 + yy)

    # 3) Plain numeric string (last resort)
    try:
        return float(s)
    except Exception:
        return None


def _summarize_yoy_metric(industry_df: pd.DataFrame,
                          company_df: pd.DataFrame,
                          metric_col: str,
                          metric_name: str,
                          grade: str) -> Optional[dict]:
    if company_df is None or company_df.empty:
        return None

    # Expect shape: company_df has ['FY', <metric_col>]
    comp = company_df.rename(columns={metric_col: "value"}).copy()

    # Try to get a sortable numeric FY; if all None, we use row order as a fallback
    comp["__fy_num__"] = comp["FY"].apply(_fy_to_num)
    if comp["__fy_num__"].notna().any():
        # Keep rows that have either a parsed FY or at least a value; we'll sort by parsed FY
        comp = comp.dropna(subset=["value"]).sort_values("__fy_num__", na_position="last")
    else:
        # Fallback: preserve incoming order and create a sequence index
        comp = comp.dropna(subset=["value"]).reset_index(drop=True)
        comp["__fy_num__"] = comp.index.astype(float)

    if comp.empty:
        return None

    # Prepare industry medians (optional)
    if industry_df is not None and not industry_df.empty:
        ind = industry_df[["FY", "p25", "p50", "p75"]].copy()
        ind["__fy_num__"] = ind["FY"].apply(_fy_to_num)
        if ind["__fy_num__"].notna().any():
            ind = ind.sort_values("__fy_num__", na_position="last")
        else:
            ind = ind.reset_index(drop=True)
            ind["__fy_num__"] = ind.index.astype(float)
    else:
        ind = pd.DataFrame(columns=["FY", "p25", "p50", "p75", "__fy_num__"])

    # Join on FY string equality; if that yields nothing at the last row, we still carry company-only
    merged = pd.merge(
        comp[["FY", "__fy_num__", "value"]],
        ind[["FY", "__fy_num__", "p25", "p50", "p75"]],
        on=["FY"], how="left", suffixes=("", "_ind")
    ).sort_values("__fy_num__")

    vals = merged["value"].tolist()
    fys  = merged["FY"].astype(str).tolist()
    n    = len(vals)

    v_first = vals[0] if n >= 1 else None
    v_last  = vals[-1] if n >= 1 else None

    def _calc_cagr(v_first: float, v_last: float, n_years: int):
        try:
            if v_first is None or v_last is None: return None
            # CAGR only meaningful on positive bases
            if v_first <= 0 or v_last <= 0 or n_years < 1: return None
            return (v_last / v_first) ** (1.0 / n_years) - 1.0
        except Exception:
            return None

    cagr = _calc_cagr(v_first, v_last, max(1, n - 1))

    last_abs_chg = last_pct_chg = None
    if n >= 2 and vals[-2] is not None and v_last is not None:
        prev = vals[-2]
        try:
            last_abs_chg = v_last - prev
            last_pct_chg = (v_last / prev - 1.0) if prev != 0 else None
        except Exception:
            last_pct_chg = None

    last_row = merged.iloc[-1]
    p25 = last_row.get("p25", None)
    p50 = last_row.get("p50", None)
    p75 = last_row.get("p75", None)

    above_median = None
    if pd.notna(p50) and v_last is not None and pd.notna(v_last):
        above_median = (v_last >= p50)

    try:
        last_bucket = _classify_bucket(v_last, p25, p50, p75, grade)
    except Exception:
        last_bucket = "—"

    show_k = 6
    ts_fy  = fys[-show_k:]
    ts_val = vals[-show_k:]
    ts_med = merged["p50"].tolist()[-show_k:] if "p50" in merged.columns else None

    return {
        "metric_name":   metric_name,
        "metric_col":    metric_col,
        "grade":         grade,
        "last_fy":       str(last_row["FY"]),
        "last_value":    v_last,
        "last_bucket":   last_bucket,
        "above_median":  above_median,
        "last_abs_chg":  last_abs_chg,
        "last_pct_chg":  last_pct_chg,
        "cagr":          cagr,
        "series_fy":     ts_fy,
        "series_value":  ts_val,
        "series_median": ts_med,
    }

def _fmt_pct(x, places=1):
    try: return f"{x*100:.{places}f}%" if x is not None and pd.notna(x) else "NA"
    except Exception: return "NA"

# =============================================================================
# GPT stubs (unchanged behavior)
# =============================================================================
def _get_openai_api_key():
    # 1. Try Streamlit secrets
    try:
        key = st.secrets["openai"]["api_key"]
        if key:
            return key
    except Exception:
        pass

    # 2. Try environment variable
    key = os.environ.get("OPENAI_API_KEY")
    if key:
        return key

    return None

def _extract_text_any(out_obj):
    try:
        # Newer Responses API convenience field
        txt = getattr(out_obj, "output_text", None)
        if txt and str(txt).strip():
            return str(txt).strip()
    except Exception:
        pass

    # Try common shapes
    try:
        # Chat Completions shape
        if hasattr(out_obj, "choices") and out_obj.choices:
            # 1) Standard chat
            msg = getattr(out_obj.choices[0], "message", None)
            if msg and getattr(msg, "content", None):
                return str(msg.content).strip()
            # 2) Responses-like content parts under choices
            content = getattr(out_obj.choices[0], "content", None)
            if isinstance(content, list) and content:
                for part in content:
                    if isinstance(part, dict) and part.get("type") in ("output_text", "text"):
                        if part.get("text"):
                            return str(part["text"]).strip()
    except Exception:
        pass

    # Responses API "output" list shape (content parts)
    try:
        output = getattr(out_obj, "output", None)
        if isinstance(output, list):
            for item in output:
                content = item.get("content") if isinstance(item, dict) else getattr(item, "content", None)
                if isinstance(content, list):
                    for part in content:
                        # part may be dict or object with attributes
                        if isinstance(part, dict):
                            if part.get("type") in ("output_text", "text") and part.get("text"):
                                return str(part["text"]).strip()
                        else:
                            if getattr(part, "type", None) in ("output_text", "text") and getattr(part, "text", None):
                                return str(getattr(part, "text")).strip()
    except Exception:
        pass

    # Final fallback: stringify object (for debugging)
    try:
        return str(out_obj).strip()
    except Exception:
        return ""

def _call_openai(system_prompt, user_prompt, api_key, model=None, max_tokens=600):
        if not api_key:
            return None, "OpenAI API key is not set."

        try:
            from openai import OpenAI
            client = OpenAI(api_key=api_key)
            mdl = (model or os.environ.get("OPENAI_MODEL") or "gpt-5").strip()

            # --- Try Responses API first ---
            try:
                out = client.responses.create(
                    model=mdl,
                    input=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user",   "content": user_prompt},
                    ],
                    # Correct for Responses API:
                    response_format={"type": "text"},         # force text section
                    reasoning={"effort": "low"},              # reduce thought token usage
                    max_output_tokens=max_tokens,
                )
                txt = _extract_text_any(out)
                if txt:
                    return txt, None
            except TypeError:
                # Retry without max_output_tokens if SDK build rejects it
                try:
                    out = client.responses.create(
                        model=mdl,
                        input=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user",   "content": user_prompt},
                        ],
                    )
                    txt = _extract_text_any(out)
                    if txt:
                        return txt, None
                except Exception:
                    pass
            except Exception:
                pass  # fall through to Chat Completions

            # --- Fallback: Chat Completions ---
            resp = client.chat.completions.create(
                model=mdl,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user",   "content": user_prompt},
                ],
                # Correct for Chat Completions:
                max_tokens=max_tokens,
            )
            txt2 = _extract_text_any(resp)
            return (txt2 if txt2 else None), None

        except Exception as e:
            return None, f"OpenAI call failed: {e}"


# --- Select the worst-performing indicator per metric type ---
def _pick_worst_per_type(summary_rows, mtype_df, max_types=None):

    # Build a map from metric column -> Type (case-insensitive, trimmed)
    col_to_type = {}
    for _, r in mtype_df.iterrows():
        c = str(r.get("Metrics_Col", "")).strip()
        t = str(r.get("Type", "")).strip() or "Uncategorized"
        if c:
            col_to_type[c] = t

    def _to_float(x):
        try:
            return float(str(x).replace(",", "").strip())
        except Exception:
            return float("nan")

    def _bucket_rank(b):
        # Needs Improvement > Satisfactory > Healthy
        order = {"Needs Improvement": 3, "Satisfactory": 2, "Healthy": 1}
        return order.get(str(b).strip(), 0)

    def _badness(row):
        grade = str(row.get("Metrics_Grade", "")).strip().lower()
        v = _to_float(row.get("value", row.get("value_str")))
        p50 = _to_float(row.get("p50", row.get("p50_str")))
        if math.isnan(v) or math.isnan(p50):
            return 0.0
        if grade == "lower":
            # higher than p50 is bad; clamp at 0 for better-than-median
            return max(v - p50, 0.0)
        else:
            # lower than p50 is bad
            return max(p50 - v, 0.0)

    # Group rows by Type
    by_type = {}
    for r in summary_rows:
        c = str(r.get("Metrics_Col", "")).strip()
        t = col_to_type.get(c, "Uncategorized")
        r = dict(r)  # shallow copy
        r["Type"] = t
        by_type.setdefault(t, []).append(r)

    # Pick the worst per Type: prefer Needs Improvement; else worst badness
    worst_per_type = []
    for t, items in by_type.items():
        # 1) try Needs Improvement first
        ni = [x for x in items if str(x.get("bucket","")) == "Needs Improvement"]
        if ni:
            cand = max(ni, key=lambda x: (_bucket_rank(x.get("bucket")), _badness(x)))
        else:
            # 2) otherwise pick the single worst by badness (ties broken by bucket rank)
            cand = max(items, key=lambda x: (_badness(x), _bucket_rank(x.get("bucket"))))
        worst_per_type.append(cand)

    # Rank the chosen types globally: most severe first
    worst_per_type.sort(key=lambda x: (_bucket_rank(x.get("bucket")), _badness(x)), reverse=True)

    # Optional cap
    if isinstance(max_types, int) and max_types > 0:
        worst_per_type = worst_per_type[:max_types]

    return worst_per_type

# --- Prompt builders ---
def _build_fin_analysis_prompt(company: str, exchange: str, industry: str, fy: str,
                                     metrics_rows: list[dict], currency: str = "") -> tuple[str, str]:

    system = (
        "You are an experienced financial analyst. Write a concise, management-ready financial analysis. "
        "Use clear, non-technical language, no jargon. Prioritize insights on fraud detection, financial analysis, and internal controls. "
        "Prioritize areas with high risk or anomalies. Do not provide investment advice."
    )

    # Compact table-like lines for the model
    lines = []
    for r in metrics_rows:
        name = r.get("Metrics_Name") or r.get("Metrics_Name")  # tolerate key shape
        val  = r.get("value_str", "NA")
        p25  = r.get("p25_str", "NA")
        p50  = r.get("p50_str", "NA")
        p75  = r.get("p75_str", "NA")
        grade= r.get("Metrics_Grade", "")
        bucket = r.get("bucket", "")
        lines.append(f"- {name}: value={val} {currency} | p25={p25} | p50={p50} | p75={p75} | grade={grade} | bucket={bucket}")

    user = (
        f"Company: {company}\nExchange: {exchange}\nIndustry: {industry}\nFY: {fy}\nCurrency: {currency}\n\n"
        "Metrics (company vs industry percentiles):\n" +
        "\n".join(lines) +
        "\n\nWrite:\n"
        "1) Executive summary (3-5 bullets) mentioning the key strengths (healthy) and key pressure points (needs improvement).\n"
        "2) Perform a root cause analysis comparing the metrics against each other and provide concise recommendations.\n"
        "For Solvency & Distress, mention about going concern probability for the company.\n"
        "Avoid acronyms unless already present in metric names. Keep to ~250 to 400 words. Bold all subheadings."
    )
    return system, user

def _build_yoy_analysis_prompt(company: str, exchange: str, industry: str,
                               summaries: list[dict], currency: str = "") -> tuple[str, str]:
    system = (
        "You are an experienced financial analyst."
        "Write a concise, management-ready analysis of YEAR-ON-YEAR trends and projected operational risks. "
        "Use clear, non-technical language (no investment advice). "
        "Prioritize metrics showing adverse trends vs industry medians or high volatility."
    )
    lines = []
    for s in summaries:
        if not s: continue
        line = (
            f"- {s['metric_name']}: "
            f"last_fy={s['last_fy']}, "
            f"last_value={('NA' if s['last_value'] is None else f'{s['last_value']:.4g}')} {currency}, "
            f"last_bucket={s['last_bucket']}, "
            f"vs_median={'above' if s['above_median'] else 'below' if s['above_median'] is not None else 'NA'}, "
            f"last_YoY={('NA' if s['last_abs_chg'] is None else f'{s['last_abs_chg']:.4g} {currency}')} "
            f"({_fmt_pct(s['last_pct_chg'])}), "
            f"CAGR={_fmt_pct(s['cagr'])}, "
            f"series_fy={','.join(map(str, s['series_fy']))}, "
            f"series_val={','.join('NA' if v is None else f'{v:.4g}' for v in s['series_value'])}"
        )
        if s.get("series_median") is not None:
            line += f", series_med={','.join('NA' if v is None else f'{v:.4g}' for v in s['series_median'])}"
        lines.append(line)

    user = (
        f"Company: {company}\nExchange: {exchange}\nIndustry: {industry}\n\n"
        "Metrics YoY snapshot (company vs industry medians):\n"
        + "\n".join(lines)
        + "\n\nWrite:\n"
        "1) Executive summary (3–5 bullets highlighting the dominant YoY trends).\n"
        "2) Projected risks and early-warning signals (tie to metrics with adverse YoY or below-median levels).\n"
        "3) Areas to monitor over the next 12 months (specific triggers and thresholds if relevant).\n"
        "Avoid acronyms unless already present in metric names. Keep to ~250 to 400 words. Bold all subheadings."
    )
    return system, user

def _build_audit_prompt(company, exchange, industry, fy, summary_rows, mtype_df,
                                       max_types=3, counts_only=False):
    system = (
        "You are an experienced internal auditor. Your expertise includes fraud detection, financial analysis, and internal controls. "
        "Given company data and industry benchmarks, identify the top 3 control areas for internal audit focus. "
        "Prioritize areas with high risk or anomalies. Avoid external audit or generic compliance steps."
        "Do not use acronyms or abbreviations in your response. Always write out the full term."
    )

    chosen = _pick_worst_per_type(summary_rows, mtype_df, max_types=max_types)
    lines = [
        f"- {r['Metrics_Name']} (grade={r['Metrics_Grade']})"
        for r in chosen]
    metrics_summary = "\n".join(lines)

    user = (
            f"Company: {company}\n"
            f"Industry: {industry}\n"
            f"Year: {fy}\n"
            f"Weakest Financial Metrics: {metrics_summary}\n\n"
            "Suggest internal control areas for audit based on potential anomalies you infer from common risk patterns in the industry that this industry."
            "For each area, include:\n"
            "1. Risk rationale.\n"
            "2. Suggested audit procedures (exceptions/fraud focus).\n"
            "3. Data required (source systems and fields).\n"
            " Bold the title of each priority area."
            "State assumptions if data is missing."
        )
    return system, user


def _convert_suggestions_to_json(readable_text: str, model: str = None) -> dict:
    api_key = _get_openai_api_key()
    if not api_key or not readable_text:
        return {}

    system_prompt = (
        "You are a meticulous converter. Convert the given auditor-readable notes into a JSON object that strictly "
        "follows the 'work_program' schema described below. Do not add fields. Do not include Markdown or prose. "
        "Only output JSON."
    )
    # Give the schema again for reliability (same keys as your mapper expects)
    schema_hint = (
        "Schema:\n"
        "{\n"
        '  "work_program": {\n'
        '    "company": "string", "exchange": "string", "industry": "string", "fy": "string",\n'
        '    "top_areas": [\n'
        '      {\n'
        '        "scope_name": "string",\n'
        '        "scope_keywords": ["string"],\n'
        '        "sub_process_name": "string",\n'
        '        "sub_process_keywords": ["string"],\n'
        '        "risk": "string",\n'
        '        "control_description": "string",\n'
        '        "procedures": ["string"],\n'
        '        "data_required": ["string"]\n'
        '      }\n'
        '    ]\n'
        "  }\n"
        "}"
    )
    user_prompt = (
        f"{schema_hint}\n\n"
        "Convert the following content into the schema. Use best-effort extraction for the three priority areas. "
        "Return ONLY JSON, no backticks:\n\n"
        f"=== INPUT START ===\n{readable_text}\n=== INPUT END ==="
    )

    text, err = _call_openai(
        system_prompt, user_prompt,
        api_key=api_key,
        model=(model or os.environ.get("OPENAI_MODEL", "gpt-5")),
        max_tokens=600
    )
    if err or not text:
        return {}

    try:
        payload = text.strip()
        # Strip any accidental fences or pre/post text
        if payload.startswith("```"):
            payload = payload.strip("`")
            payload = payload[payload.find("{"): payload.rfind("}") + 1]
        return json.loads(payload)
    except Exception:
        return {}


# =============================================================================
# Utility (numeric validation)
# =============================================================================

def _reset_app_state(full_cache_reset: bool = False):

    # 1) Clear known keys from this app
    keys_to_clear = [
        # Autofill scaffolding
        "_autofill_company", "_autofill_fy", "_autofill_exchange",

        # Sidebar inputs
        "company_name_input", "current_assets", "current_liabilities", "inventory",
        "operating_cf", "capex", "revenue", "ebitda", "cost_of_revenue",

        # Selection + data snapshots
        "selection", "pct_wide", "pct_tidy",
        "company_name", "exchange", "industry", "fy",
        "mtype_bm", "mtype_yoy",

        # Tab 1 analysis
        "fin_analysis_text", "fin_analysis_pdf_bytes",

        # (If you added charts export later)
        "bm_figs", "bm_figs_key",

        # Tab 3 suggestions
        "ai_audit_suggestions", "ai_audit_suggestions_pdf_bytes",
        "ai_work_program", "ai_mappings",

        # Tab 5 observations
        "audit_observations",
    ]

    for k in keys_to_clear:
        st.session_state.pop(k, None)

    # 2) Clear any dynamic keys created by file uploaders and chart PDF cache
    #    (uploaders in Tab 4 use keys like "doc_<scope>_<sub>_<i>")
    for k in list(st.session_state.keys()):
        if k.startswith("doc_") or k.startswith("bm_charts_pdf_bytes"):
            st.session_state.pop(k, None)

    # 3) Remove temp uploads folder (created in Tab 4)
    try:
        shutil.rmtree("uploads", ignore_errors=True)
    except Exception:
        pass

    # 4) Clear Tab 1 analysis via your existing helper (if present)
    try:
        _clear_tab1_analysis()
    except Exception:
        pass

    # 5) Optionally clear cached data (expensive; forces reload of Excel, etc.)
    if full_cache_reset:
        try:
            st.cache_data.clear()
        except Exception:
            pass

def _is_number_str(x: Optional[str]) -> bool:
    try:
        sx = "" if x is None else str(x).strip()
        if sx == "":
            return False
        float(sx)
        return True
    except Exception:
        return False

# ---- Normalization & tokenization ----
def _norm_text(s: str | None) -> str:
    if s is None:
        return ""
    s = str(s).lower()
    s = re.sub(r"[^a-z0-9\s/\-+&]", " ", s)  # keep some useful symbols
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _tokens(s: str) -> set[str]:
    return set(_norm_text(s).split())

def _bigrams(s: str) -> set[tuple[str, str]]:
    toks = _norm_text(s).split()
    return set(zip(toks, toks[1:])) if len(toks) > 1 else set()

def _jaccard(a: set, b: set) -> float:
    if not a and not b:
        return 0.0
    return len(a & b) / max(1, len(a | b))

def _seq_ratio(a: str, b: str) -> float:
    return SequenceMatcher(None, _norm_text(a), _norm_text(b)).ratio()

# ---- Optional synonym map from an Excel "Synonyms" sheet (Type, From, To) ----
@st.cache_data(show_spinner=False)
def load_synonyms_map(path: str = AUDIT_DB_FILE) -> dict[tuple[str, str], str]:
    try:
        xl = pd.ExcelFile(path, engine="openpyxl")
        if "Synonyms" not in xl.sheet_names:
            return {}
        df = xl.parse("Synonyms")
        # expected columns: Type (Scope/Sub-process), From, To
        out = {}
        for _, r in df.iterrows():
            typ = str(r.get("Type", "")).strip()
            frm = _norm_text(r.get("From", ""))
            to  = str(r.get("To", "")).strip()
            if typ and frm and to:
                out[(typ, frm)] = to
        return out
    except Exception:
        return {}

# ---- Scoring (label + optional keywords vs candidate + context) ----
def _score_label_to_candidate(label: str, candidate: str,
                              extra_keywords: list[str] | None = None,
                              candidate_context: list[str] | None = None) -> float:
    ek = " ".join(extra_keywords or [])
    comp_a = f"{label} {ek}".strip()
    ctx = " ".join(candidate_context or [])
    # base similarities
    s1 = _seq_ratio(comp_a, candidate)
    j1 = _jaccard(_tokens(comp_a), _tokens(candidate))
    j2 = _jaccard(_bigrams(comp_a), _bigrams(candidate))
    # optional context boost
    j_ctx = _jaccard(_tokens(comp_a), _tokens(ctx)) * 0.5 if ctx else 0.0
    # weighted blend (tune to taste)
    score = 0.50 * s1 + 0.30 * j1 + 0.15 * j2 + 0.05 * j_ctx
    return float(score)

# ---- Scope candidate ranking ----
def _rank_scopes(area: dict, audit_df: pd.DataFrame, synonyms: dict | None = None) -> list[tuple[str, float]]:
    label = area.get("scope_name") or area.get("scope") or ""
    kws   = area.get("scope_keywords", []) or []
    if synonyms:
        label = synonyms.get(("Scope", _norm_text(label)), label)

    scopes = sorted(audit_df["Scope"].dropna().unique().tolist())
    # Context per scope: sub-process names + a sample of risk/controls
    ctx_by_scope = {}
    for s in scopes:
        subnames = audit_df.loc[audit_df["Scope"] == s, "Sub-process"].astype(str).unique().tolist()
        risks    = audit_df.loc[audit_df["Scope"] == s, "Risk"].astype(str).tolist()[:10]
        ctrls    = audit_df.loc[audit_df["Scope"] == s, "Control Description"].astype(str).tolist()[:10]
        ctx_by_scope[s] = ["; ".join(subnames), "; ".join(risks), "; ".join(ctrls)]

    scored = []
    for s in scopes:
        scored.append((s, _score_label_to_candidate(label, s, kws, candidate_context=ctx_by_scope[s])))
    scored.sort(key=lambda x: x[1], reverse=True)
    return scored

# ---- Sub-process candidate ranking (within a scope) ----
def _rank_subprocess(area: dict, audit_df: pd.DataFrame, scope: str,
                     synonyms: dict | None = None) -> list[tuple[str, float]]:
    label = area.get("sub_process_name") or area.get("sub_process") or ""
    kws   = area.get("sub_process_keywords", []) or []
    if synonyms:
        label = synonyms.get(("Sub-process", _norm_text(label)), label)

    subs = sorted(audit_df.loc[audit_df["Scope"] == scope, "Sub-process"].dropna().unique().tolist())
    ctx_by_sub = {}
    for sp in subs:
        r = audit_df[(audit_df["Scope"] == scope) & (audit_df["Sub-process"] == sp)]
        risk  = r["Risk"].iloc[0] if not r.empty else ""
        ctrl  = r["Control Description"].iloc[0] if not r.empty else ""
        ctx_by_sub[sp] = [str(risk), str(ctrl)]

    scored = []
    for sp in subs:
        scored.append((sp, _score_label_to_candidate(label, sp, kws, candidate_context=ctx_by_sub[sp])))
    scored.sort(key=lambda x: x[1], reverse=True)
    return scored

# ---- Full mapping of all top_areas ----
def map_ai_to_master(ai_json: dict, audit_df: pd.DataFrame,
                     auto_threshold: float = 0.82, review_threshold: float = 0.60,
                     synonyms_map: dict | None = None) -> list[dict]:
    areas = (ai_json or {}).get("work_program", {}).get("top_areas", [])
    mappings = []
    for a in areas:
        scope_cands = _rank_scopes(a, audit_df, synonyms_map)
        scope_match, scope_score = scope_cands[0]
        if scope_score < review_threshold:
            mappings.append({
                "area": a,
                "status": "no_match",
                "reason": "scope_below_threshold",
                "scope_candidates": scope_cands[:5],
            })
            continue
        sub_cands = _rank_subprocess(a, audit_df, scope_match, synonyms_map)
        sub_match, sub_score = sub_cands[0]
        status = "auto" if (scope_score >= auto_threshold and sub_score >= auto_threshold) else "review"

        row = audit_df[(audit_df["Scope"] == scope_match) & (audit_df["Sub-process"] == sub_match)].iloc[0]
        mappings.append({
            "area": a,
            "status": status,
            "mapped_scope": scope_match,
            "mapped_sub_process": sub_match,
            "scope_score": float(scope_score),
            "sub_score": float(sub_score),
            "row_index": int(row.name),
            "scope_candidates": scope_cands[:5],
            "sub_candidates": sub_cands[:5],
        })
    return mappings

# =============================================================================
# MAIN UI
# =============================================================================
def main():
    st.set_page_config(page_title="Industry Benchmark Analysis Dashboard", layout="wide")
    st.markdown(CHIP_CSS, unsafe_allow_html=True)
    st.title("Industry Benchmark Analysis Dashboard")

    # Load data
    sheets = load_all_sheets()
    data_df = sheets[DATA_SHEET]
    mtype_df = sheets[METRICS_TYPE_SHEET]

    # Validate metrics_type sheet
    required_cols = ["Type", "Metrics_Name", "Metrics_Col", "Metrics_Description", "Metrics_Grade"]
    for c in required_cols:
        if c not in mtype_df.columns:
            st.error(f"'{METRICS_TYPE_SHEET}' sheet missing column: {c}")
            st.stop()

    # Prepare metrics and percentiles
    metric_cols = infer_metric_columns(data_df)
    data_df = to_numeric(data_df, metric_cols)
    pct_wide, pct_tidy = compute_percentiles(data_df, metric_cols)

    # --- Sidebar: Always-visible Company Input (with auto-population) ---
    
    st.sidebar.button(
        "↻ Refresh / Clear All",
        type="secondary",
        use_container_width=True,
        on_click=lambda: _reset_app_state(full_cache_reset=True)
    )

    st.sidebar.header("Company Input")
    
    # Exchange list (for initial default)
    exch_list = sorted(data_df["EXCHANGE"].dropna().astype(str).unique().tolist())
    default_exch_idx = exch_list.index("SGX") if "SGX" in exch_list else 0

    # Session scaffolding for autofill
    ss = st.session_state
    ss.setdefault("_autofill_company", "")
    ss.setdefault("_autofill_fy", "")
    ss.setdefault("_autofill_exchange", "")
    for k in [
        "current_assets", "current_liabilities", "inventory",
        "operating_cf", "capex", "revenue", "ebitda", "cost_of_revenue"
    ]:
        ss.setdefault(k, "")

    # 1) Company Name (first)
    typed_company = st.sidebar.text_input("Company Name", key="company_name_input")

    # Suggestions (up to 10)
    suggestions = get_company_matches(data_df, typed_company) if typed_company else []
    if typed_company and suggestions and typed_company.strip() not in suggestions:
        st.sidebar.caption("Suggestions: " + ", ".join(suggestions))

    # Find exact company block (case-insensitive)
    df_company = find_exact_company(data_df, typed_company) if typed_company else pd.DataFrame()

    # 2) Exchange (defaults to company’s exchange if exact match, else SGX or first)
    default_exchange = (
        pick_best_profile(df_company).get("EXCHANGE", exch_list[default_exch_idx]) if not df_company.empty
        else exch_list[default_exch_idx]
    )
    exchange = st.sidebar.selectbox("Exchange", exch_list, index=exch_list.index(default_exchange))

    # 3) Industry list depends on selected exchange
    ind_list = sorted(
        data_df.loc[data_df["EXCHANGE"].astype(str) == exchange, "INDUSTRY"].dropna().astype(str).unique().tolist()
    )
    profile_industry = pick_best_profile(df_company).get("INDUSTRY", "") if not df_company.empty else ""
    default_industry = profile_industry if profile_industry in ind_list else (ind_list[0] if ind_list else "")
    industry = st.sidebar.selectbox(
        "Industry",
        ind_list,
        index=(ind_list.index(default_industry) if default_industry in ind_list else 0)
    )

    # 4) Financial Year list
    if not df_company.empty:
        years_arr = sorted(df_company["FY"].dropna().astype(str).unique().tolist())
    else:
        cand = data_df[(data_df["EXCHANGE"].astype(str) == exchange) & (data_df["INDUSTRY"].astype(str) == industry)]
        years_arr = sorted(cand["FY"].dropna().astype(str).unique().tolist())
    fy = st.sidebar.selectbox("Financial Year", years_arr, index=(len(years_arr) - 1 if years_arr else 0), key="fy_select")

    # Currency label by exchange
    ccy = CURRENCY_BY_EXCHANGE.get(exchange, "")

    # Auto-populate financial inputs when (company, fy, exchange) changes and exact match exists
    autofill_trigger = f"{typed_company or ''}|{fy or ''}|{exchange or ''}"
    last_trigger = f"{ss.get('_autofill_company','')}|{ss.get('_autofill_fy','')}|{ss.get('_autofill_exchange','')}"
    if not df_company.empty and fy and autofill_trigger != last_trigger:
        snapshot_vals = read_key_financials_for(df_company, prefer_fy=fy)  # <-- already scaled to units here

        mapping = [
            ("Current Assets",       "current_assets"),
            ("Current Liabilities",  "current_liabilities"),
            ("Inventory",            "inventory"),
            ("Operating Cash Flow",  "operating_cf"),
            ("Capital Expenditure",  "capex"),
            ("Revenue",              "revenue"),
            ("EBITDA",               "ebitda"),
            ("Cost of Revenue",      "cost_of_revenue"),
        ]
        for label, key in mapping:
            if label in snapshot_vals:
                ss[key] = f"{snapshot_vals[label]}"
            else:
                ss.setdefault(key, ss.get(key, ""))

        ss["_autofill_company"]  = typed_company or ""
        ss["_autofill_fy"]       = fy or ""
        ss["_autofill_exchange"] = exchange or ""

    # 5) Key financials — ALWAYS visible text inputs (editable; auto-populated when matched)
    current_assets      = st.sidebar.text_input(f"Current Assets ({ccy})", key="current_assets")
    current_liabilities = st.sidebar.text_input(f"Current Liabilities ({ccy})", key="current_liabilities")
    inventory           = st.sidebar.text_input(f"Inventory ({ccy})", key="inventory")
    operating_cf        = st.sidebar.text_input(f"Operating Cash Flow ({ccy})", key="operating_cf")
    capex               = st.sidebar.text_input(f"Capital Expenditure ({ccy})", key="capex")
    revenue             = st.sidebar.text_input(f"Revenue ({ccy})", key="revenue")
    ebitda              = st.sidebar.text_input(f"EBITDA ({ccy})", key="ebitda")
    cost_of_revenue     = st.sidebar.text_input(f"Cost of Revenue ({ccy})", key="cost_of_revenue")

    # 6) Numeric gating (like your earlier prototype)
    numeric_keys = [
        "current_assets", "current_liabilities", "inventory", "operating_cf",
        "capex", "revenue", "ebitda", "cost_of_revenue"
    ]
    numeric_ok = all(_is_number_str(ss.get(k)) for k in numeric_keys)
    identifiers_ok = bool(exchange and industry and fy and typed_company and typed_company.strip())

    if not numeric_ok:
        st.sidebar.warning("Please complete all financial fields with valid numbers before submitting.")

    # --- Footer ---
    st.sidebar.markdown("<hr>", unsafe_allow_html=True)
    st.sidebar.caption("version 3.5 | 2026")

    # Submit is disabled until all eight financial fields are numeric and identifiers are set
    can_submit = identifiers_ok and numeric_ok
    submit = st.sidebar.button("Submit", type="primary", disabled=not can_submit)

    # --- after numeric gating and when user clicks Submit ---
    if not submit and "selection" not in st.session_state:
        st.info("Fill in the Company Input on the left. All financial fields must be numeric. Click **Submit** to load benchmarks.")
        st.stop()

    if submit:
        st.session_state["selection"] = {
            "company": typed_company.strip(),
            "exchange": exchange,
            "industry": industry,
            "fy": str(fy)  # store as string for consistency
        }
        st.session_state["pct_wide"] = pct_wide
        st.session_state["pct_tidy"] = pct_tidy
        st.session_state["mtype_bm"] = None   # current selection for Benchmarking tab
        st.session_state["mtype_yoy"] = None  # current selection for YoY tab

    sel = st.session_state["selection"]
    company = sel["company"]; exch = sel["exchange"]; ind = sel["industry"]; fy_sel = sel["fy"]
    pct_wide = st.session_state["pct_wide"]; pct_tidy = st.session_state["pct_tidy"]

    # -------------------------------------------------------------------------
    # Body content (after Submit): show Metrics Type first, then benchmarking UI
    # -------------------------------------------------------------------------

    # Freeze choices into session (for consistency across tabs)
    ss["company_name"] = typed_company
    ss["exchange"] = exchange
    ss["industry"] = industry
    ss["fy"] = fy
    ss["pct_wide"] = pct_wide
    ss["pct_tidy"] = pct_tidy

    company = ss["company_name"]
    exch = ss["exchange"]
    ind = ss["industry"]
    fy_sel = ss["fy"]

    st.markdown(f"**Company:** {company} &nbsp;&nbsp; **Exchange:** {exch} &nbsp;&nbsp; "
                f"**Industry:** {ind} &nbsp;&nbsp; **FY:** {fy_sel}")
        
    tab_bm, tab_yoy, tab_audit, tab_wp, tab_obs = st.tabs(["1.Benchmarking (Selected FY)", "2.YoY Trend", "3.Suggested Top 3 Audit Areas",
        "4.Audit Work Program", "5.Observations Summary"
    ])

    # -------------------------------------------------------------------------
    # TAB 1 — Benchmarking (Selected FY)
    # -------------------------------------------------------------------------

    with tab_bm:
        st.markdown(
            '<div class="legend">Legend: '
            '<span class="chip chip-green">Healthy</span>'
            '<span class="chip chip-amber">Satisfactory</span>'
            '<span class="chip chip-red">Needs Improvement</span></div><div class="hr"></div>',
            unsafe_allow_html=True
        )

        type_list = list(mtype_df["Type"].dropna().astype(str).unique())
        # Safer default index handling (avoids ValueError if session value not in list)
        if st.session_state.get("mtype_bm") in type_list:
            default_idx = type_list.index(st.session_state["mtype_bm"])
        else:
            default_idx = 0

        mtype = st.selectbox("Select a metrics type", type_list, index=default_idx, key="mtype_bm", on_change=_clear_tab1_analysis)
        subset = mtype_df[mtype_df["Type"] == mtype].copy()
        st.caption(f"Showing metrics for type: **{mtype}** ({len(subset)} metrics)")

        comp_slice = slice_company_row_for_fy(data_df, exch, ind, fy_sel, company)
        p_row = _try_get_percentile_row(pct_wide, exch, ind, fy_sel)

        grid = st.columns(2)
        assembled_for_llm = []
        skipped_bm = []
        bm_figs = []

        for i, (_, rec) in enumerate(subset.iterrows()):
            m_col   = str(rec["Metrics_Col"]).strip()
            m_name  = str(rec["Metrics_Name"]).strip()
            m_desc  = str(rec["Metrics_Description"]).strip()
            m_grade = str(rec["Metrics_Grade"]).strip()

            val = pd.to_numeric(comp_slice.iloc[0].get(m_col, np.nan), errors="coerce") if not comp_slice.empty else np.nan
            p25 = p50 = p75 = np.nan
            if p_row is not None and (m_col, "p25") in getattr(p_row, "index", []):
                p25 = pd.to_numeric(p_row[(m_col, "p25")], errors="coerce")
                p50 = pd.to_numeric(p_row[(m_col, "p50")], errors="coerce")
                p75 = pd.to_numeric(p_row[(m_col, "p75")], errors="coerce")

            # Skip if nothing to show 
            if (pd.isna(val) and pd.isna(p25) and pd.isna(p50) and pd.isna(p75)):
                skipped_bm.append(m_name)
                continue

            bucket = _classify_bucket(val, p25, p50, p75, m_grade)

            with grid[i % 2]:
                with st.container():
                    st.markdown(f"**{m_name}**")
                    st.markdown(f'<span class="chip {_chip_class(bucket)}">{bucket}</span>', unsafe_allow_html=True)
                    st.caption(m_desc)

                    fig = _plot_benchmark(
                        val, p25, p50, p75, m_grade,
                        company_name=company,
                        band_thickness=0.16
                    )
                    st.plotly_chart(fig, use_container_width=True, key=f"bm_{m_col}")
                    _fmt = lambda x: "NA" if pd.isna(x) else f"{x:.4g}"
                    st.caption(f"Company={_fmt(val)} · p25={_fmt(p25)} · p50={_fmt(p50)} · p75={_fmt(p75)}")
                    
                    # Keep figure for PDF bundling
                    bm_figs.append((m_name, bucket, fig))

            # keep your LLM assembly
            assembled_for_llm.append({
                "Metrics_Name": m_name, "Metrics_Col": m_col, "Metrics_Grade": m_grade,
                "value": None if pd.isna(val) else float(val), "value_str": ("NA" if pd.isna(val) else f"{val:.4g}"),
                "p25": None if pd.isna(p25) else float(p25), "p50": None if pd.isna(p50) else float(p50),
                "p75": None if pd.isna(p75) else float(p75),
                "p25_str": ("NA" if pd.isna(p25) else f"{p25:.4g}"),
                "p50_str": ("NA" if pd.isna(p50) else f"{p50:.4g}"),
                "p75_str": ("NA" if pd.isna(p75) else f"{p75:.4g}"),
                "bucket": bucket,
            })

        if skipped_bm:
            with st.expander("Skipped metrics (not available for this FY / industry / company)", expanded=False):
                st.write(", ".join(skipped_bm))
        ss["bm_figs"] = bm_figs
        
        st.divider()
        st.subheader("Financial Metrics - Analysis")

        fin_model = st.text_input("Model", os.environ.get("OPENAI_MODEL", "gpt-5"), key="bm_model")

        currency_label = ccy if 'ccy' in locals() else ""

        btn_fin = st.button(
            "Generate Analysis",
            type="primary",
            key="btn_fin_analysis",
            disabled=(len(assembled_for_llm) == 0)
        )

        if btn_fin:
            api_key_fin = _get_openai_api_key()
            if not api_key_fin:
                st.error("OpenAI API key is missing. Set it in Streamlit secrets or OPENAI_API_KEY.")
            else:
                sys_p, usr_p = _build_fin_analysis_prompt(
                    company=company, exchange=exch, industry=ind, fy=str(fy_sel),
                    metrics_rows=assembled_for_llm, currency=currency_label
                )
                with st.spinner("Calling model and drafting financial analysis..."):
                    fin_text, fin_err = _call_openai(
                        sys_p, usr_p, api_key=api_key_fin, model=fin_model, max_tokens=600
                    )
                if fin_err:
                    st.error(f"API Error: {fin_err}")
                elif fin_text:
                    st.session_state["fin_analysis_text"] = fin_text.strip()
                    st.success("Financial analysis generated.")
                else:
                    st.warning("No output received from the model. Try a smaller prompt or different model.")

        if st.session_state.get("fin_analysis_text"):
            text = st.session_state["fin_analysis_text"]
            st.markdown(text)

            # Generate & cache PDF bytes (optional caching)
            if st.session_state.get("fin_analysis_text") and st.session_state.get("bm_figs"):
                text = st.session_state["fin_analysis_text"]
                figs = st.session_state["bm_figs"]

                # a cache key that changes when inputs change
                combo_key = f"{company}|{exch}|{ind}|{fy_sel}|{mtype}|{len(figs)}|{hash(text)}"
                if ss.get("bm_fin_combo_key") != combo_key:
                    ss.pop("bm_fin_combo_html_bytes", None)
                    ss["bm_fin_combo_key"] = combo_key


                if "bm_fin_combo_html_bytes" not in ss:
                    ss["bm_fin_combo_html_bytes"] = analysis_and_charts_to_html_bytes(
                        analysis_md_text=text,
                        figs=figs,
                        title=f"{company} — Analysis & Benchmarking ({fy_sel})",
                        subtitle=f"{exch} · {ind}",
                        skipped=skipped_bm,
                    )

                st.download_button(
                    "Download analysis + charts (.html)",
                    data=ss["bm_fin_combo_html_bytes"],
                    file_name="analysis_and_benchmarking.html",
                    mime="text/html",
                    use_container_width=True,
                    key="dl_fin_plus_charts_html",
                )

    # -------------------------------------------------------------------------
    # TAB 2 — YoY Trend
    # -------------------------------------------------------------------------

    with tab_yoy:
        # --- Header / selectors ---
        type_list2 = list(mtype_df["Type"].dropna().astype(str).unique().tolist())
        if st.session_state.get("mtype_yoy") in type_list2:
            default_idx2 = type_list2.index(st.session_state["mtype_yoy"])
        else:
            default_idx2 = 0
        mtype2 = st.selectbox("Select a metrics type for YoY trend", type_list2, key="mtype_yoy", index=default_idx2)

        subset2 = mtype_df[mtype_df["Type"] == mtype2].copy()
        st.caption(f"Industry percentiles by FY for: **{exch} · {ind}**")

        # --- Data slices (frozen by sidebar Submit) ---
        tidy = pct_tidy[
            (pct_tidy["EXCHANGE"].astype(str) == exch) & (pct_tidy["INDUSTRY"].astype(str) == ind)
        ]
        company_series = yoy_company_series(data_df, exch, ind, company)

        # --- Layout & accumulators ---
        grid = st.columns(2)
        skipped_yoy: list[str] = []
        yoy_figs: list[tuple[str, str, go.Figure]] = []   # (metric_name, bucket_last, fig)
        yoy_summaries: list[dict] = []                    # compact LLM summaries per metric

        # --- Per-metric charts + collections ---
        for i, (_, rec) in enumerate(subset2.iterrows()):
            m_col = str(rec["Metrics_Col"]).strip()
            m_name = str(rec["Metrics_Name"]).strip()
            m_grade = str(rec["Metrics_Grade"]).strip()

            # Industry percentiles for YoY (may be empty)
            ind_df = (
                tidy[tidy["metric"] == m_col][["FY", "p25", "p50", "p75"]]
                .dropna()
                .sort_values("FY")
            )

            # Company series (safe take); skip if the metric column isn't present for this company
            comp_df, missing = _safe_take(company_series, ["FY", m_col])
            if missing:
                skipped_yoy.append(m_name)
                continue
            comp_df = comp_df.dropna().sort_values("FY")

            # If both company and industry are empty, skip
            if (ind_df is None or ind_df.empty) and (comp_df is None or comp_df.empty):
                skipped_yoy.append(m_name)
                continue

            # Render chart
            with grid[i % 2]:
                st.markdown(f"**{m_name}**")
                fig = _plot_yoy(ind_df, comp_df, m_col, m_grade)
                st.plotly_chart(fig, use_container_width=True, key=f"yoy_{m_col}")

            # Compute last-year bucket (for HTML chip label)
            bucket_last = "—"
            try:
                if not comp_df.empty:
                    last_fy = comp_df["FY"].iloc[-1]
                    last_val = comp_df[m_col].iloc[-1]
                    if not ind_df.empty and (last_fy in ind_df["FY"].values):
                        pr = ind_df[ind_df["FY"] == last_fy].iloc[0]
                        bucket_last = _classify_bucket(
                            float(last_val),
                            float(pr.get("p25", float("nan"))),
                            float(pr.get("p50", float("nan"))),
                            float(pr.get("p75", float("nan"))),
                            m_grade
                        )
            except Exception:
                bucket_last = "—"

            # >>> SAVE chart for export
            yoy_figs.append((m_name, bucket_last, fig))

            # >>> Build compact YoY summary for GPT (if helper is present)

            if "_summarize_yoy_metric" in globals():
                summary = _summarize_yoy_metric(
                    industry_df=ind_df,
                    company_df=comp_df[["FY", m_col]].rename(columns={m_col: m_col}),
                    metric_col=m_col,
                    metric_name=m_name,
                    grade=m_grade
                )
                if summary:
                    yoy_summaries.append(summary)
            else:
                st.error("YoY helper `_summarize_yoy_metric` is missing. Please add it before generating analysis.")

        # Show skipped list, if any
        if skipped_yoy:
            with st.expander("Skipped YoY metrics (missing in company data / industry percentiles)", expanded=False):
                st.write(", ".join(skipped_yoy))

        # Persist figures in session for any later use
        st.session_state["yoy_figs"] = yoy_figs

        # ============================
        # YoY ANALYSIS (GPT write-up)
        # ============================
        st.divider()
        st.subheader("YoY Trend - Analysis")
        yoy_model = st.text_input("Model", os.environ.get("OPENAI_MODEL", "gpt-5"), key="yoy_model")

        btn_yoy = st.button(
            "Generate YoY Analysis",
            type="primary",
            key="btn_yoy_analysis",
            disabled=False  # allow click; we'll guard below if summaries are empty
        )

        if btn_yoy:
            if "_build_yoy_analysis_prompt" not in globals():
                st.error("YoY helper `_build_yoy_analysis_prompt` is missing. Add it above `main()`.")
            elif "_summarize_yoy_metric" not in globals():
                st.error("YoY helper `_summarize_yoy_metric` is missing. Add it above `main()`.")
            else:
                # 1) Proceed only if you actually have at least one valid summary
                valid_summaries = [s for s in yoy_summaries if s]
                if not valid_summaries:
                    # You reached this branch because the helpers exist but produced no summaries.
                    # Typical causes: company series empty after dropna(), name mismatch, or only industry series.
                    st.warning(
                        "No YoY summaries were built (company series may be empty/sparse). "
                        "Verify the company name matches `ENTITY_NAME` exactly and that the metric has numeric values across ≥ 2 FY."
                    )
                else:
                    # 2) Build the proper YoY prompt and call the model
                    api_key_fin = _get_openai_api_key()
                    if not api_key_fin:
                        st.error("OpenAI API key is missing. Set it in Streamlit secrets or OPENAI_API_KEY.")
                    else:
                        currency_label = CURRENCY_BY_EXCHANGE.get(exch, "")
                        sys_p, usr_p = _build_yoy_analysis_prompt(
                            company=company, exchange=exch, industry=ind,
                            summaries=valid_summaries, currency=currency_label
                        )
                        with st.spinner("Calling model and drafting YoY analysis..."):
                            yoy_text, yoy_err = _call_openai(
                                sys_p, usr_p, api_key=api_key_fin, model=yoy_model, max_tokens=600
                            )
                        if yoy_err:
                            st.error(f"API Error: {yoy_err}")
                        elif yoy_text:
                            st.session_state["yoy_analysis_text"] = str(yoy_text).strip()
                            st.success("YoY analysis generated.")
                        else:
                            st.warning("No output received from the model. Try a smaller prompt or different model.")


        # Show analysis if present
        if st.session_state.get("yoy_analysis_text"):
            st.markdown(st.session_state["yoy_analysis_text"])

        # ============================
        # HTML EXPORT (charts + write-up)
        # ============================
        has_charts = bool(yoy_figs)
        if has_charts:
            # Determine the FY window shown
            try:
                fy_list = company_series["FY"].astype(str).tolist()
                period = f"{min(fy_list)}–{max(fy_list)}" if fy_list else str(fy_sel)
            except Exception:
                period = str(fy_sel)

            yoy_text = (st.session_state.get("yoy_analysis_text") or "").strip()

            # Cache key to avoid regenerating the same HTML
            combo_key = f"{company}|{exch}|{ind}|{mtype2}|{len(yoy_figs)}|{hash(yoy_text)}"
            if st.session_state.get("yoy_combo_key") != combo_key:
                st.session_state.pop("yoy_combo_html_bytes", None)
                st.session_state["yoy_combo_key"] = combo_key

            if "yoy_combo_html_bytes" not in st.session_state:
                st.session_state["yoy_combo_html_bytes"] = analysis_and_charts_to_html_bytes(
                    analysis_md_text=yoy_text,  # empty string allowed (charts-only export)
                    figs=yoy_figs,
                    title=f"{company} — YoY Trends & Projected Risks ({period})",
                    subtitle=f"{exch} · {ind}",
                    skipped=skipped_yoy,
                )

            btn_label = "Download YoY analysis + charts (.html)" if yoy_text else "Download YoY charts (.html)"
            st.download_button(
                btn_label,
                data=st.session_state["yoy_combo_html_bytes"],
                file_name="yoy_trends_and_risks.html",
                mime="text/html",
                use_container_width=True,
                key="dl_yoy_plus_charts_html",
            )
        else:
            st.info("No YoY charts to export yet.")


    # -------------------------------------------------------------------------
    # TAB 3 — Suggested Audit Areas (Top 3)
    # -------------------------------------------------------------------------
    with tab_audit: 
        st.subheader("Suggested Top 3 Audit Areas")
        st.caption("Analyzes selected company financial metrics, industry benchmarks and suggests the top 3 key auditable areas for assessment.")
        
        # --- 1. DATA FILTERING (Fixes the NameError) ---
        sel_mask = (
            (data_df["EXCHANGE"].astype(str) == exch)
            & (data_df["INDUSTRY"].astype(str) == ind)
            & (data_df["FY"].astype(str) == str(fy_sel))
        )
        df_slice = data_df.loc[sel_mask]
        
        # Identify the specific company row
        comp_row = df_slice[df_slice["ENTITY_NAME"].astype(str).str.strip().str.lower() == company.strip().lower()]
        if comp_row.empty and not df_slice.empty:
            comp_row = df_slice.iloc[[0]]
            
        # Identify the benchmark/percentile row
        p_row = _try_get_percentile_row(pct_wide, exch, ind, str(fy_sel))

        # --- 2. PREPARE METRICS FOR THE HELPER ---
        all_metrics_to_rank = []
        if not comp_row.empty:
            for _, r in mtype_df.iterrows():
                c = str(r["Metrics_Col"]).strip()
                n = str(r["Metrics_Name"]).strip()
                g = str(r["Metrics_Grade"]).strip()

                val = pd.to_numeric(comp_row.iloc[0].get(c, np.nan), errors="coerce")
                if pd.isna(val):
                    continue

                p25 = p50 = p75 = np.nan
                if p_row is not None and (c, "p25") in p_row.index:
                    p25 = pd.to_numeric(p_row[(c, "p25")], errors="coerce")
                    p50 = pd.to_numeric(p_row[(c, "p50")], errors="coerce")
                    p75 = pd.to_numeric(p_row[(c, "p75")], errors="coerce")

                bucket = _classify_bucket(val, p25, p50, p75, g)
                
                all_metrics_to_rank.append({
                    "Metrics_Name": n,
                    "Metrics_Col": c,
                    "Metrics_Grade": g,
                    "value": float(val),
                    "value_str": f"{val:.4g}",
                    "p50": float(p50) if pd.notna(p50) else None,
                    "bucket": bucket,
                })

        # --- 3. UI AND GENERATION ---
        model = st.text_input("Input Model :", os.environ.get("OPENAI_MODEL", "gpt-5"), key="audit_model")
        generate = st.button("Generate Audit Suggestions", type="primary")

        if generate:
            if not all_metrics_to_rank:
                st.warning("No problematic metrics found to analyze.")
            else:
                system_prompt, user_prompt = _build_audit_prompt(
                    company, exch, ind, str(fy_sel),
                    all_metrics_to_rank, mtype_df, max_types=3
                )
                api_key_for_call = _get_openai_api_key()
                if not api_key_for_call:
                    st.error("OpenAI API key is missing.")
                else:
                    with st.spinner("Calling model and analyzing risk areas..."):
                        text, err = _call_openai(
                            system_prompt, user_prompt,
                            api_key=api_key_for_call, model=model, max_tokens=600
                        )

                    if err:
                        st.error(f"API Error: {err}")
                    elif text:
                        # 1) Show readable results on-page
                        st.success("Audit suggestions generated!")
                        st.session_state["ai_audit_suggestions"] = text

            if st.session_state.get("ai_audit_suggestions"):
                st.markdown(st.session_state["ai_audit_suggestions"])
                # Build once and cache in session
                if "ai_audit_suggestions_pdf_bytes" not in st.session_state:
                    st.session_state["ai_audit_suggestions_pdf_bytes"] = md_to_pdf_bytes(
                        st.session_state["ai_audit_suggestions"],
                        title=f"{company} – Suggested Audit Areas ({fy_sel})",
                        author="Auto-generated by the dashboard",
                        subheaders={
                            "Company": company,                          
                            "Exchange": exch,                             
                            "Industry": ind,                             
                            "FY": str(fy_sel),                            
                        })

                st.download_button(
                    "Download Suggestions (.pdf)",
                    data=st.session_state["ai_audit_suggestions_pdf_bytes"],
                    file_name="audit_suggestions.pdf",
                    mime="application/pdf",
                    key="dl_audit_suggestions_pdf"
                )

    # -------------------------------------------------------------------------
    # TAB 4 — Audit Work Program
    # -------------------------------------------------------------------------

    with tab_wp:
        st.subheader("Audit Work Program — AI-Assisted Audit Testing")

        # Gate: do not show info until Tab 3 has generated content
        has_suggestions = bool(st.session_state.get("ai_audit_suggestions"))
        if not has_suggestions:
            st.info("No suggested audit areas yet. Generate suggested audit areas in Tab 3.")
        else:
            # Load the full master (no dependency on Tab 3)
            audit_df = load_audit_db()
            scopes, subs_by_scope = audit_vocab(audit_df)

            # Guard: no scopes available
            if not scopes:
                st.error("No scopes found in the Audit Work Program master.")
                st.stop()

            # Pick safe defaults
            default_scope = st.session_state.get("wp_scope", scopes[0])
            if default_scope not in scopes:
                default_scope = scopes[0]

            # Scope selector (full list)
            sel_scope = st.selectbox(
                "1) Scope",
                scopes,
                index=scopes.index(default_scope),
                key="wp_scope"
            )

            # Resolve sub-process list for the selected scope
            sub_list = subs_by_scope.get(sel_scope, [])
            if not sub_list:
                st.warning("No sub-processes found for the selected scope. Please choose another scope.")
                st.stop()

            # Default sub-process: last used if still valid, else first
            default_sub = st.session_state.get("wp_subproc", sub_list[0])
            if default_sub not in sub_list:
                default_sub = sub_list[0]

            # Sub-process selector (full list for the selected scope)
            sel_sub = st.selectbox(
                "2) Sub-process",
                sub_list,
                index=sub_list.index(default_sub),
                key="wp_subproc"
            )

            # Retrieve row and show Risk/Control 
            row = audit_df[(audit_df["Scope"] == sel_scope) & (audit_df["Sub-process"] == sel_sub)]
            if row.empty:
                st.error("No work program row found for this selection.")
                st.stop()

            rec = row.iloc[0]
            st.markdown("**Risk**")
            st.info(rec["Risk"])
            st.markdown("**Control Description**")
            st.info(rec["Control Description"])

            # --- Documents required (column F) -> Yes/No + conditional uploaders
            st.markdown("**Documents required**")
            doc_items = _parse_documents_required(rec["Documents required"])

            answers = {}  # { label: {"have": "Yes"/"No", "file": <UploadedFile or None>} }
            have_keys = []
            for i, lab in enumerate(doc_items):
                col_q, col_u = st.columns([1, 1.4])

                have_key = f"have_{sel_scope}_{sel_sub}_{i}"
                have_val = col_q.radio(
                    f"Do you have this document?\n\n**{lab}**",
                    options=["No", "Yes"], index=0, horizontal=True, key=have_key
                )
                have_keys.append(have_key)

                file_obj = None
                if have_val == "Yes":
                    file_obj = col_u.file_uploader(
                        f"Upload: {lab}", type=None, accept_multiple_files=False,
                        key=f"doc_{sel_scope}_{sel_sub}_{i}"
                    )
                else:
                    # Ensure any previous file is cleared if user flips Yes -> No
                    st.session_state.pop(f"doc_{sel_scope}_{sel_sub}_{i}", None)

                answers[lab] = {"have": have_val, "file": file_obj}

            # Validation: all questions answered AND all "Yes" have a file
            # (radios always have a value, default "No", so 'answered' is implicit)
            docs_ok = all(
                (ans["have"] == "No") or (ans["have"] == "Yes" and ans["file"] is not None)
                for ans in answers.values()
            )

            if not docs_ok:
                st.warning("Please answer **Yes/No** for every item. For each **Yes**, upload the document to proceed.")
            st.markdown("---")


            # ---- OpenAI Draft Runner: open files, extract excerpts, evaluate vs test steps ----
            
            submit_wp = st.button(
                "Run Audit Test Steps & Draft Observations", type="primary", key=f"run_{sel_scope}_{sel_sub}",disabled=not docs_ok)

            if submit_wp:
                # 1) Save + parse files
                os.makedirs("uploads", exist_ok=True)
                saved = []
                parsed_docs = []

                for label, ans in answers.items():
                    file = ans.get("file")
                    if ans.get("have") == "Yes" and file is not None:
                        safe_name = f"{int(time.time())}_{company}_{sel_scope}_{sel_sub}_{os.path.basename(file.name)}".replace(" ", "_")
                        path = os.path.join("uploads", safe_name)
                        with open(path, "wb") as f:
                            f.write(file.getbuffer())

                        text, meta = extract_text_from_upload(file)
                        meta["label"] = label
                        parsed_docs.append({"text": text, "meta": meta})
                        saved.append({"label": label, "path": path, "original_name": file.name, "type": meta.get("type", "")})

                # 2) Build concise evidence pack
                evidence = build_relevant_excerpts(rec["Audit Test Steps"], parsed_docs)
                ev_lines = []
                for j, ev in enumerate(evidence, start=1):
                    fname = ev["name"]
                    ev_lines.append(f"[snippet {j}] file={fname} · score={ev['score']:.2f}\n{ev['excerpt']}\n")
                evidence_block = "\n".join(ev_lines) if ev_lines else "(no snippets extracted)"

                # 3) Call OpenAI (reuses your existing helpers and model setting)
                model = os.environ.get("OPENAI_MODEL", st.session_state.get("audit_model", "gpt-5"))
                api_key = _get_openai_api_key()
                if not api_key:
                    st.error("OpenAI API key is missing. Set it in Streamlit secrets or OPENAI_API_KEY.")
                    st.stop()

                system_prompt = (
                    "You are an experienced internal auditor. You will be given:\n"
                    "1) Company context (scope, sub-process, risk, control);\n"
                    "2) Audit Test Steps; and\n"
                    "3) Evidence excerpts extracted from uploaded documents.\n\n"
                    "Task:\n"
                    "- Evaluate the evidence against each audit step.\n"
                    "- Identify potential observations (exceptions, control gaps, anomalies or missing evidence).\n"
                    "- Classify severity (High / Medium / Low).\n"
                    "- Propose a concise root cause and recommendation.\n"
                    "- Where applicable, cite the evidence by file name and “snippet#”.\n\n"
                    "Return ONLY valid JSON of the shape:\n"
                    "{\n"
                    "  \"observations\": [\n"
                    "    {\"observation\":\"...\", \"severity\":\"High|Medium|Low\", \"root_cause\":\"...\", \"recommendation\":\"...\", \"evidence_refs\":[{\"file\":\"...\",\"snippet_id\":1}]}\n"
                    "  ]\n"
                    "}\n"
                    "If there are no issues, return {\"observations\": []}."
                )

                user_prompt = (
                    f"Company: {company}\n"
                    f"Scope: {sel_scope}\n"
                    f"Sub-process: {sel_sub}\n\n"
                    f"Risk:\n{rec['Risk']}\n\n"
                    f"Control Description:\n{rec['Control Description']}\n\n"
                    f"Audit Test Steps:\n{rec['Audit Test Steps']}\n\n"
                    f"Evidence Excerpts (top-ranked):\n{evidence_block}\n\n"
                    "Draft observations (if any). Use only the JSON schema specified."
                )

                with st.spinner("Calling model to analyze evidence and draft observations..."):
                    text, err = _call_openai(system_prompt, user_prompt, api_key=api_key, model=model, max_tokens=600)

                # 4) Parse JSON, store for Tab 5
                observations = []
                if err:
                    st.error(f"OpenAI call failed: {err}")
                else:
                    try:
                        payload = text.strip()
                        if payload.startswith("```"):
                            payload = payload.strip("`")
                            payload = payload[payload.find("{"): payload.rfind("}")+1]
                        out = json.loads(payload)
                        observations = out.get("observations", [])
                        if not isinstance(observations, list):
                            observations = []
                    except Exception as e:
                        st.warning(f"Could not parse model JSON: {e}")
                        observations = []

                if observations:
                    st.success("Observations drafted. Refer to 5 - Observations Summary.")
                    if "audit_observations" not in st.session_state:
                        st.session_state["audit_observations"] = []
                    # Map snippet_id -> original file short name; then to saved path
                    snippet_to_file = {}
                    for j, ev in enumerate(evidence, start=1):
                        snippet_to_file[j] = ev["name"]

                    enriched = []
                    for o in observations:
                        refs = o.get("evidence_refs", []) or []
                        files_from_refs = []
                        for r in refs:
                            sn = int(r.get("snippet_id", 0))
                            fname = snippet_to_file.get(sn)
                            if fname:
                                match = next((s for s in saved if s["original_name"] == fname), None)
                                files_from_refs.append(match["path"] if match else fname)
                        if not files_from_refs:
                            files_from_refs = [s["path"] for s in saved]  # fallback

                        enriched.append({
                            "company": company, "exchange": exch, "industry": ind, "fy": str(fy_sel),
                            "scope": sel_scope, "sub_process": sel_sub,
                            "observation": o.get("observation", ""),
                            "severity": o.get("severity", ""),
                            "root_cause": o.get("root_cause", ""),
                            "recommendation": o.get("recommendation", ""),
                            "evidence_links": files_from_refs,
                        })
                    st.session_state["audit_observations"].extend(enriched)
                else:
                    st.info("No issues drafted by the model; consider adding more evidence or refining steps.")


    # -------------------------------------------------------------------------
    # TAB 5 — Observations
    # -------------------------------------------------------------------------

    with tab_obs:
        st.subheader("Audit Observations Summary")

        obs = st.session_state.get("audit_observations", "")
        if not obs:
            st.info("No observations yet. Run a work program in Tab 4.")
        else:
            df_obs = pd.DataFrame(obs)

            # ---- Column order + friendly display names ----
            col_order = [
                "company", "exchange", "industry", "fy",
                "scope", "sub_process",
                "severity", "observation", "root_cause", "recommendation",
                "evidence_links",
            ]
            # Keep only columns that exist, in preferred order
            col_order = [c for c in col_order if c in df_obs.columns]
            df_obs = df_obs[col_order].copy()

            # Friendly headers map
            headers = {
                "company": "Company",
                "exchange": "Exchange",
                "industry": "Industry",
                "fy": "FY",
                "scope": "Scope",
                "sub_process": "Sub-process",
                "severity": "Severity",
                "observation": "Observation",
                "root_cause": "Root cause",
                "recommendation": "Recommendation",
                "evidence_links": "Evidence",
            }
            df_obs = df_obs.rename(columns=headers)

            # ---- Show short evidence names in the grid (export keeps full paths) ----
            def _shorten(paths):
                try:
                    if isinstance(paths, (list, tuple)):
                        import os
                        return ", ".join(os.path.basename(p) for p in paths)
                    return paths
                except Exception:
                    return paths

            df_display = df_obs.copy()
            if "Evidence" in df_display.columns:
                df_display["Evidence"] = df_display["Evidence"].apply(_shorten)

            # ---- Add CSS to enable wrapping and improve grid readability ----
            st.markdown(
                """
                <style>
                /* Make dataframe cells wrap and grow vertically */
                div[data-testid="stDataFrame"] div[role="gridcell"] {
                    white-space: normal !important;
                    overflow-wrap: anywhere !important;
                    line-height: 1.25rem;        /* 20px for readability */
                    padding-top: 6px; 
                    padding-bottom: 6px;
                }
                /* Header wrap as well (for smaller screens) */
                div[data-testid="stDataFrame"] div[role="columnheader"] {
                    white-space: normal !important;
                    overflow-wrap: anywhere !important;
                }
                /* Slightly taller rows for readability */
                div[data-testid="stDataFrame"] div[role="row"] {
                    align-items: start;
                }
                </style>
                """,
                unsafe_allow_html=True,
            )

            # ---- Render as a responsive grid with sensible column widths ----
            from streamlit import column_config as cc

            st.dataframe(
                df_display,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Company": cc.TextColumn(width="small"),
                    "Exchange": cc.TextColumn(width="small"),
                    "Industry": cc.TextColumn(width="small"),
                    "FY": cc.TextColumn(width="small"),
                    "Scope": cc.TextColumn(width="medium"),
                    "Sub-process": cc.TextColumn(width="medium"),
                    "Severity": cc.TextColumn(width="small"),
                    # Long text columns get larger widths (wrapping enabled via CSS above)
                    "Observation": cc.TextColumn(width="large"),
                    "Root cause": cc.TextColumn(width="large"),
                    "Recommendation": cc.TextColumn(width="large"),
                    "Evidence": cc.TextColumn(width="medium"),
                },
                height=520,  
            )

            # ---- Keep Excel export with full paths (no shortening) ----
            bio = BytesIO()
            with pd.ExcelWriter(bio, engine="openpyxl") as xw:
                # Export original df_obs with full evidence_links values
                df_obs.to_excel(xw, index=False, sheet_name="Observations")
            st.download_button(
                "Download Excel",
                data=bio.getvalue(),
                file_name="audit_observations.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

# =============================================================================
# Entrypoint
# =============================================================================
if __name__ == "__main__":
    _spawn_streamlit()

